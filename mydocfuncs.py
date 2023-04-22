import docx

from myconstants import *

TEST_READ_FILE = r'C:\Test_folder\ТестоваяСтатья.docx'
TEST_WRITE_FILE = r'C:\Test_folder\TestNew.docx'


def print_doc(doc_4_print):
    text = []
    for paragraph in doc_4_print.paragraphs:
        text.append(paragraph.text)
    print('\n'.join(text))


def update_paragraph(paragraph, new_text):
    '''
    Затирает в параграфе все пробеги (runs), а весь текст объединяет в одном (первом) пробеге.
    :param paragraph: Ссылка на параграф.
    :param new_text: Замещаемый текст.
    :return:
    '''
    for one_run in paragraph.runs:
        one_run.text = ''

    paragraph.runs[0].text = new_text

    return paragraph


def get_all_refs(doc_object, p_only_in_text=False):
    '''
    Возвращает список всех кодов (##код), встречающихся в тексте.
    :param doc_object: Загруженный документ WinWord.
    :param p_only_in_text: Если True, то собираются коды только из текста - игнорируются коды в списке литературы.
    :return: Список кодов.
    '''
    # Если передан параметр p_only_in_text равный True,
    # то проверяется начало строки и если в самом начале
    # стоит две "решётки", то строка игнорируется!
    found_refs_in_text = []

    # Обработаем каждый параграф
    for one_paragraph in doc_object.paragraphs:
        paragraph_text = one_paragraph.text
        start_poz = paragraph_text.find(REFS_START_CHAR_IN_TEXT)
        if p_only_in_text and start_poz == 0:
            pass
        else:
            while start_poz != -1:
                end_poz = 0
                paragraph_text = paragraph_text[start_poz:]
                if len(paragraph_text) >= (len(REFS_START_CHAR_IN_TEXT) + 2):
                    for poz, selected_char in enumerate(paragraph_text):
                        if selected_char in REFS_STOP_CHARS:
                            end_poz = poz
                            break

                if end_poz != 0:
                    new_mark = paragraph_text[0:end_poz]
                    if new_mark not in found_refs_in_text:
                        found_refs_in_text.append(new_mark)

                paragraph_text = paragraph_text[(len(new_mark) + 1):]

                start_poz = paragraph_text.find(REFS_START_CHAR_IN_TEXT)

    return found_refs_in_text


def find_refs_in_list(doc_object):
    '''
    Обрабатывает список литературы, который обычно располагается в конце документа.
    :param doc_object: Загруженный документ WinWord
    :return: Собранный список списков, содержащий информацию о литературе, содержащий:
    1) код конкретного элемента из списка литературы (начинается с ##);
    2) физический номер в списке в исходном загруженном тексте;
    3) текстовое описание источника литературы, включая код.
    '''

    found_refs = []

    # Обработаем каждый параграф
    for num_paragraph, one_paragraph in enumerate(doc_object.paragraphs):
        paragraph_text = one_paragraph.text.strip()
        if len(paragraph_text) > 2:
            if paragraph_text[0:len(REFS_START_CHAR)] == REFS_START_CHAR and not paragraph_text[len(REFS_START_CHAR)] in REFS_STOP_CHARS:
                for one_char in REFS_STOP_CHARS:
                    paragraph_text = paragraph_text.split(one_char)[0]

                # После завершения цикла в переменной paragraph_text остаётся
                # ссылка на этот параграф типа ##001, которая записывается на первом месте в списке:
                found_refs.append([paragraph_text, num_paragraph, one_paragraph.text.strip()])

    return found_refs


def replace_refs_in_doc(doc_object, old_ref, new_ref):
    '''
    Меняет код ссылки на литературу, но её вычисленный номер.
    :param doc_object: Загруженный документа WinWord.
    :param old_ref: Старая ссылка (код = ##код)
    :param new_ref: Вычисленный номер в спске литературы.
    :return:
    '''
    for one_paragraph in doc_object.paragraphs:
        collect_ref_parts_in_one_run_and_replace(one_paragraph, old_ref, new_ref)
        for one_run in one_paragraph.runs:
            one_run.text = one_run.text.replace(old_ref, new_ref)


def collect_ref_parts_in_one_run_and_replace(one_paragraph, old_ref, new_ref):
    while old_ref in one_paragraph.text:
        text = one_paragraph.text
        start_idx = text.find(old_ref)
        end_idx = start_idx + len(old_ref) - 1
        char_idx_counter = 0
        for num_run, one_run in enumerate(one_paragraph.runs):
            this_run_text = one_run.text
            for one_char in this_run_text:
                char_idx_counter += 1
                if char_idx_counter < start_idx + 1:
                    # Это ещё не начло ссылки
                    continue
                if char_idx_counter == start_idx + 1:
                    # Это самое начало ссылки!
                    # Запомним в каком пробеге она началась:
                    begin_ref_num_run = num_run
                elif char_idx_counter <= end_idx + 1 and num_run == begin_ref_num_run:
                    # Это продолжение ссылки. Пробег тот же. Ничего не делаем.
                    continue
                elif char_idx_counter <= end_idx + 1 and num_run != begin_ref_num_run:
                    # Это продолжение ссылки, которое попало в другой пробег.
                    # Её надо перенести в правильный пробег.
                    # 1) Перенесём символ в правильный пробег:
                    one_paragraph.runs[begin_ref_num_run].text = one_paragraph.runs[begin_ref_num_run].text + one_char
                    # 2) Уберём этот символ из начала "неправильного" пробега:
                    one_paragraph.runs[num_run].text = one_paragraph.runs[num_run].text[1:]
                else:
                    # Всё ссылка закончилась.
                    # Заменим ссылку на новую и переходим к обработке следующей.
                    one_paragraph.runs[begin_ref_num_run].text = one_paragraph.runs[begin_ref_num_run].text.replace(old_ref, new_ref)
                    break


def replace_ref_paragraphs(ui, doc_object, ordered_refs, refs_list, all_unused_refs):
    '''
    Сортировка списка литературы в той очерёдности, в которой встречаются ссылки в тексте документа.
    :param ui: Ссылка на MainWindow.ui
    :param doc_object: Ссылка на загруженный документа WinWord
    :param ordered_refs: Список ссылок, встречающихся в тексте, для которых есть номер в списке литературы.
    :param refs_list: Список с информацией о списке литературы: [##код, позиция в списке, текст]
    :param all_unused_refs: все не используемые коды из списка литературы.
    :return: Список текстовых описаний списка литературы.
    '''

    ui.progressBar2.setMinimum(0)
    ui.progressBar2.setMaximum(len(ordered_refs) + (len(all_unused_refs) * ui.del_unused_refs.isChecked()))
    left_delimiter = ui.befor_num.text()
    right_delimiter = ui.after_num.text()

    refs_result = []
    # Переберём все параграфы из списка литературы и расположим
    # их в том порядке в каком на них встречаются ссылки в тексте:
    # Пройдемся по ПОЛНОМУ отсортированному списку кодов:
    for num_ref, one_ref in enumerate(ordered_refs):
        pFlag = (ui.del_unused_refs.isChecked() and one_ref in all_unused_refs)
        paragraph_poz = refs_list[num_ref][1]
        ui.progressBar2.setValue(ui.progressBar2.value() + 1)

        for one_paragraph in refs_list:
            if one_paragraph[0] == one_ref:
                new_paragraph_text = one_paragraph[2]

                if not pFlag:
                    new_paragraph_text = f"{left_delimiter}{num_ref + 1}{right_delimiter} {new_paragraph_text[(len(one_ref)):].strip()}"

                update_paragraph(doc_object.paragraphs[paragraph_poz], new_paragraph_text)

                if not pFlag:
                    replace_refs_in_doc(doc_object, one_ref, f"{num_ref + 1}")
                    refs_result.append(new_paragraph_text)

                break

    # Если установлен признак "удалить неиспользуемые ссылки", то обработаем их
    if ui.del_unused_refs.isChecked():
        for num_ref, one_ref in enumerate(all_unused_refs):
            ui.progressBar2.setValue(ui.progressBar2.value() + 1)
            for one_paragraph in doc_object.paragraphs:
                if one_paragraph.text.find(one_ref) == 0:
                    p = one_paragraph._p
                    parent_element = p.getparent()
                    parent_element.remove(p)

                    break

    return refs_result


def get_refs_errors(doc_object, ordered_refs):
    errors = []
    for one_ref in ordered_refs:
        for one_paragraph in doc_object.paragraphs:
            if one_ref in one_paragraph.text:
                if one_ref not in errors:
                    errors.append(one_ref)

    return errors


def get_docx_object(file_name):
    return docx.Document(file_name)


def save_docx_object(doc_object, file_name):
    doc_object.save(file_name)


if __name__ == "__main__":
    doc = get_docx_object(TEST_READ_FILE)
    all_ordered_refs = get_all_refs(doc)
    print(all_ordered_refs)

    all_refs_list = find_refs_in_list(doc)
    print(all_refs_list)

    replace_ref_paragraphs(doc, all_ordered_refs, all_refs_list)

    save_docx_object(doc, TEST_WRITE_FILE)
