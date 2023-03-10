# main.py

import os
import sys
import win32com.client
import pandas as pd
import numpy as np
import datetime as dt
import time
import warnings
import docx

from constants import *

TEST_READ_FILE = r'C:\Test_folder\ТестоваяСтатья.docx'
TEST_READ_FILE = r'C:\Test_folder\Test.docx'
TEST_WRITE_FILE = r'C:\Test_folder\TestNew.docx'


def print_doc(doc_4_print):
    text = []
    for paragraph in doc_4_print.paragraphs:
        text.append(paragraph.text)
    print('\n'.join(text))


def update_paragraph(paragraph, new_text):
    for one_run in paragraph.runs:
        one_run.text = ''

    paragraph.runs[0].text = new_text

    return paragraph


def get_all_refs_in_text(doc_object):
    found_refs_in_text = []
    for one_paragraph in doc_object.paragraphs:
        paragraph_text = one_paragraph.text
        start_poz = paragraph_text.find(REFS_START_CHAR_IN_TEXT)
        end_poz = 0
        if start_poz > 0:
            paragraph_text = paragraph_text[(start_poz + 1):]
            if len(paragraph_text) >= 3:
                for poz, selected_char in enumerate(paragraph_text):
                    if selected_char in REFS_STOP_CHARS:
                        end_poz = poz
                        break

            if end_poz != 0:
                new_mark = paragraph_text[0:end_poz]
                if new_mark not in found_refs_in_text:
                    found_refs_in_text.append(new_mark)

    return found_refs_in_text


def find_refs_list(doc_object):
    found_refs = []
    for num_paragraph, one_paragraph in enumerate(doc_object.paragraphs):
        paragraph_text = one_paragraph.text.strip()
        if len(paragraph_text) > 2:
            if paragraph_text[0] == REFS_START_CHAR and not paragraph_text[1] in REFS_STOP_CHARS:
                for one_char in REFS_STOP_CHARS:
                    paragraph_text = paragraph_text.split(one_char)[0]

                found_refs.append([paragraph_text, num_paragraph, one_paragraph.text.strip()])

    return found_refs


def replace_refs_in_text(doc_object, old_ref, new_ref):
    for one_paragraph in doc_object.paragraphs:
        for one_run in one_paragraph.runs:
            one_run.text = one_run.text.replace(old_ref, new_ref)

    return doc_object


def replace_ref_paragraphs(doc_object, ordered_refs, refs_list):
    for num_ref, one_ref in enumerate(ordered_refs):
        paragraph_poz = refs_list[num_ref][1]
        for one_paragraph in refs_list:
            if one_paragraph[0] == one_ref:
                new_paragraph_text = one_paragraph[2]
                new_paragraph_text = f"{num_ref + 1}. {new_paragraph_text[(len(one_ref)):].strip()}"
                update_paragraph(doc_object.paragraphs[paragraph_poz], new_paragraph_text)
                doc_object = replace_refs_in_text(doc_object, one_ref, f"{num_ref + 1}")
                break

    return doc_object


if __name__ == "__main__":
    doc = docx.Document(TEST_READ_FILE)
    all_ordered_refs = get_all_refs_in_text(doc)
    print(all_ordered_refs)

    all_refs_list = find_refs_list(doc)
    print(all_refs_list)

    doc = replace_ref_paragraphs(doc, all_ordered_refs, all_refs_list)

    doc.save(TEST_WRITE_FILE)

